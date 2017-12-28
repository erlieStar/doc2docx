#导入界面库
from docx2excel import Ui_MainWindow
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox, QAction
import sys
#导入工具库
from docx import Document
import xlwt
import xlrd
from xlrd import book
from xlutils.copy import copy
import sys
import os
import time
import shutil

#docx2excel的逻辑部分


#要录入的Excel表的地址
startExcel = ''
#最后生成的excel，这个库只能保存xls格式的文件
endExcel = ''
#包含装换信息的文件，主要包括正确的文件，格式错误的文件，不是简单表，中途运行出错
msgExcel = ''
#保存格式错误的文件夹
formatErrorDir = ''
#保存中途错误的文件夹
runErrorDir = ''

# sheet1的字典
dict1 = {}
# sheet2的字典
dict2 = {}
# sheet3的字典
dict3 = {}
# sheet4的字典
dict4 = {}
# sheet5的字典
dict5 = {}

successList = []
errorList = []
notSimpleList = []
runErrorList = []
totalList = []

re = book
we = xlwt.Workbook()


#判断是否是英文
def isEnglish(checkStr):
    for ch in checkStr:
        if u'\u4e00' <= ch <= u'\u9fff':
            return False
    return True

def readSubTemplate(table, start, end, index):

    columnList = table.columns
    columnLength = len(columnList)

    dict = {}

    for rowIndex in range(start,end):
        str = ''
        sum = 0
        total = 0
        for columnIndex in range(columnLength):
            cell = table.cell(rowIndex,columnIndex)
            if columnIndex == 0:
                str = cell.text
                sum = columnIndex
                total = 1
                continue
            if cell.text == str:
                sum += columnIndex
                total += 1
            else:
                tempIndex = int(sum * 1.0 / total)
                dict.setdefault(str,[rowIndex,tempIndex])
                sum = columnIndex
                total = 1
                str = cell.text
        tempIndex = int(sum * 1.0 / total)
        dict.setdefault(str, [rowIndex, tempIndex])

    global dict2, dict3, dict4, dict5
    if index == 2:
        dict2 = dict
    elif index == 3:
        dict3 = dict
    elif index == 4:
        dict4 = dict
    elif index == 5:
        dict5 = dict

#读取模板表
def readTemplate(dialog):

    str = dialog.wordTemLineEdit.text().strip();
    document = Document(str)
    tempTable = document.tables
    table = tempTable[0]

    columnList = table.columns
    columnLength = len(columnList)

    for rowIndex in range(2,8):
        str = ''
        sum = 0
        total = 0
        flag = True
        for columnIndex in range(columnLength):
            cell = table.cell(rowIndex,columnIndex)
            if not isEnglish(cell.text):
                continue
            if flag:
                str = cell.text
                sum = columnIndex
                total = 1
                flag = False
                continue
            if cell.text == str:
                sum += columnIndex
                total += 1
            else:
                tempIndex = int(sum * 1.0 / total)
                dict1.setdefault(str, [rowIndex, tempIndex])
                sum = columnIndex
                total = 1
                str = cell.text
        tempIndex = int(sum * 1.0 / total)
        dict1.setdefault(str, [rowIndex, tempIndex])

    #生产工段/工艺信息
    readSubTemplate(table, 11, 12, 2)
    #锅炉信息
    readSubTemplate(table, 22, 23, 3)
    #露天堆场信息
    readSubTemplate(table, 28, 29, 4)
    #有机溶剂使用信息
    readSubTemplate(table, 16, 17, 5)


# 写第一页的sheet
def writeFirstSheet(wordName, document, table, sheetIndex, row):

    # print ("写第一个sheet页面")
    sheet = we.get_sheet(sheetIndex - 1)
    sheet1dict = {}
    for key in dict1:
        tempList = dict1[key]
        for index in range(0, 1):
            x = tempList[index]
            y = tempList[index + 1]
            sheet1dict.setdefault(key, table.cell(x, y).text)

    # 读取填表人和手机号
    str = ''
    paList = document.paragraphs
    for index in range(len(paList)):
        tempStr = paList[index].text
        str += tempStr

    if '：' in str:
        list = str.split('：')
        tempStr1 = list[1].replace('手机号', '').strip()
        tempStr2 = list[2].replace('填表时间', '').strip()
        sheet1dict.setdefault('informant', tempStr1)
        sheet1dict.setdefault('tel', tempStr2)

    sheet1dict.setdefault('companyId', wordName)

    tempSheet = re.sheet_by_index(sheetIndex - 1)
    list1 = tempSheet.row_values(0)
    for excelIndex in range(len(list1)):
        for key in sheet1dict:
            if list1[excelIndex] == key:
                sheet.write(row, excelIndex, sheet1dict[key])

firstRow = 7
sheetRow = [7, 7, 7, 7]

# 写第二页以后的sheet
def writeSheet(wordName, table, sheetIndex):

    # print ("写第" + str(sheetIndex) + "个sheet页面")
    global dict2, dict3, dict4, dict5
    sheetdict = {}
    dict = {}
    if sheetIndex == 2:
        dict = dict2
    elif sheetIndex == 3:
        dict = dict3
    elif sheetIndex == 4:
        dict = dict4
    elif sheetIndex == 5:
        dict = dict5

    max1 = 0
    for key in dict:
        if key == 'equipId':
            continue
        tempList = dict[key]
        for index in range(0, 1):
            x = tempList[index]
            y = tempList[index + 1]
            tempList2 = []

            data = table.cell(x, y).text
            if data != '':
                tempList2.append(data)

            data = table.cell(x + 1, y).text
            if data != '':
                if len(tempList2) == 1:
                    tempList2.append(data)
                else:
                    tempList2.append('')
                    tempList2.append(data)

            data = table.cell(x + 2, y).text
            if data != '':
                if len(tempList2) == 2:
                    tempList2.append(data)
                else:
                    tempList2.append('')
                    tempList2.append('')
                    tempList2.append(data)
            sheetdict.setdefault(key, tempList2)
            if len(tempList2) > max1:
                max1 = len(tempList2)

    if max1 == 0:
        max1 = max1 + 1

    # 填写所属企业编号
    sheetdict.setdefault('companyId', wordName)

    # 写编号信息
    if sheetIndex != 5:
        tempList3 = dict['equipId']
        tempList3 = []
        for index in range(0, max1):
            tempList3.append(index + 1)
            sheetdict.setdefault('equipId', tempList3)

    sheet = we.get_sheet(sheetIndex - 1)
    tempSheet = re.sheet_by_index(sheetIndex - 1)
    # print(sheetIndex, tempSheet.name)
    list = tempSheet.row_values(0)
    for excelIndex in range(len(list)):
        for key in sheetdict:
            if list[excelIndex] == key:
                tempSheet2Total = sheetRow[sheetIndex - 2]
                if key == 'companyId':
                    for index in range(tempSheet2Total, tempSheet2Total + max1):
                        sheet.write(index, excelIndex, sheetdict[key])
                else:
                    tempList = sheetdict[key]
                    # 所在的行数
                    for x in tempList:
                        sheet.write(tempSheet2Total, excelIndex, x)
                        tempSheet2Total += 1

    sheetRow[sheetIndex - 2] = sheetRow[sheetIndex - 2] + max1

# 测试文件是否缺行
def checkFile(table):

    x1 = table.cell(8, 0).text
    if not '生产工段/工艺信息' in x1:
        return False

    x1 = table.cell(14, 0).text
    if not '有机溶剂使用信息' in x1:
        return False

    x1 = table.cell(19, 0).text
    if not '锅炉信息' in x1:
        return False

    x1 = table.cell(25, 0).text
    if not '露天堆场信息' in x1:
        return False

    rowList = table.rows

    if len(rowList) <= 31:
        return False
    x1 = table.cell(31, 0).text
    if not '备注' in x1:
        return False

    return True

# 将word中数据写入excel
def writeExcel(wordName, wordUrl, dialog):

    dialog.textEdit.append("写入" + wordName)
    document = Document(wordUrl)
    tempTable = document.tables

    if len(tempTable) != 1:
        notSimpleList.append(wordName)
        print('不是简单表------->' + wordName)
        return
    table = tempTable[0]

    # 检查表格的格式是否正确
    if not checkFile(table):
        print('文件错误------->' + wordName)
        errorList.append(wordName)
        fileName = wordUrl.split("/")[-1]
        shutil.copyfile(wordUrl,formatErrorDir + '/' + fileName)
        return

    global firstRow
    row = firstRow
    writeFirstSheet(wordName, document, table, 1, row)
    firstRow += 1

    writeSheet(wordName, table, 2)
    writeSheet(wordName, table, 3)
    writeSheet(wordName, table, 4)
    writeSheet(wordName, table, 5)

    we.save(endExcel)

    successList.append(wordName)

# 将转换信息写入excel
def writeMsg(dialog):

    dialog.textEdit.append("生成（转换信息.xls）")
    excel = xlwt.Workbook(encoding='utf-8')
    # 这个是指定sheet页的名称
    sheet1 = excel.add_sheet('统计信息')
    sheet2 = excel.add_sheet('详细信息')
    sheet1.write(0, 0, '文件总数')
    sheet1.write(0, 1, '录入正确')
    sheet1.write(0, 2, '格式错误')
    sheet1.write(0, 3, '不是简表')
    sheet1.write(0, 4, '中途出错')

    sheet1.write(1, 0, len(totalList))
    sheet1.write(1, 1, len(successList))
    sheet1.write(1, 2, len(errorList))
    sheet1.write(1, 3, len(notSimpleList))
    sheet1.write(1, 4, len(runErrorList))

    sheet2.write(0, 0, '文件总数')
    sheet2.write(0, 1, '录入正确')
    sheet2.write(0, 2, '格式错误')
    sheet2.write(0, 3, '不是简表')
    sheet2.write(0, 4, '中途出错')

    row = 1
    for x in successList:
        sheet2.write(row, 0, x)
        row += 1

    row = 1
    for x in successList:
        sheet2.write(row, 1, x)
        row += 1

    row = 1
    for x in errorList:
        sheet2.write(row, 2, x)
        row += 1

    row = 1
    for x in notSimpleList:
        sheet2.write(row, 3, x)
        row += 1

    row = 1
    for x in runErrorList:
        sheet2.write(row, 4, x)
        row += 1

    excel.save(msgExcel)
    dialog.statusbar.showMessage("录入结束，请到Excel表格所在文件夹查看（生成的表.xls，转换信息表.xls）")
    dialog.setOp(True)


class SimpleDialogForm(Ui_MainWindow, QMainWindow):

    def __init__(self, parent = None):
        super(SimpleDialogForm, self).__init__()
        self.setupUi(self)#在此设置界面

        #父类的progressBar的值为24,这里设置为0
        self.progressBar.setProperty("value", 0)

        #新建关于的条目
        self.about = QAction("关于")
        self.contact = QAction("联系")
        #加上帮助菜单栏
        helpMenu = self.menubar.addMenu("帮助")
        #帮助菜单栏上加上条目
        helpMenu.addAction(self.about)
        helpMenu.addAction(self.contact)
        #选中word目录绑定的槽函数
        self.wordDirBtn.clicked.connect(self.setWordDirUrl)
        #选中word模板表绑定的槽函数
        self.wordTemBtn.clicked.connect(self.setWordTemUrl)
        #选中excel表绑定的槽函数
        self.excelBtn.clicked.connect(self.setExcelUrl)
        #开始按钮绑定的槽函数
        self.startBtn.clicked.connect(self.startTrans)
        #输入框文本内容发生改变时，初始化图形界面
        self.wordDirLineEdit.textChanged.connect(self.initGUI)
        self.wordTemLineEdit.textChanged.connect(self.initGUI)
        self.excelLineEdit.textChanged.connect(self.initGUI)
        #显示关于信息栏
        self.about.triggered.connect(self.showAbout)
        #显示联系作者信息栏
        self.contact.triggered.connect(self.contactAuthor)

    def showAbout(self):
        msgBox = QMessageBox(QMessageBox.Question, "关于", "选中Word目录，Word模板表，Excel表格，点击开始转换按钮\n"
                                                         "转换完成后在Excel文件夹下有（转换信息.xls）,记录了录入的\n"
                                                         "文件总数,录入正确的文件数,录入错误的文件数等详细信息，如\n"
                                                         "出现录入出现偏移的情况，是模板表和要转换表格式不一样，\n"
                                                         "建议每次录入前，先测试一下模板表是否正确")
        msgBox.exec()


    def contactAuthor(self):
        msgBox = QMessageBox(QMessageBox.Information, "联系作者", "如您有更好的意见或建议，请联系作者QQ:290059742")
        msgBox.exec()

    def setWordDirUrl(self):
        str = QFileDialog.getExistingDirectory(self,"打开Word所在目录",r"C:\Users\Administrator\Desktop")
        self.wordDirLineEdit.setText(str)

    def setWordTemUrl(self):
        str, ok  = QFileDialog.getOpenFileName(self, "打开Word模板表",r"C:\Users\Administrator\Desktop","Word File(*.docx)")
        if ok:
            self.wordTemLineEdit.setText(str)


    def setExcelUrl(self):
        str, ok = QFileDialog.getOpenFileName(self, "打开Excel表格",r"C:\Users\Administrator\Desktop","Excel File(*.xls *.xlsx)")
        if ok:
            self.excelLineEdit.setText(str)

    #初始化配置，即确定（生成的表.xls和转换信息表.xls）所在的位置
    def initConfig(self):
        global endExcel, msgExcel
        # 在模板excel所在的目录下生成（生成的表.xls）和（转换信息表.xls）
        str = self.excelLineEdit.text().strip()
        name = []
        if '/' in str:
            name = str.split('/')
        elif '\\' in str:
            name = str.split('\\')
        tempExcel = ''
        for index in range(len(name) - 1):
            tempExcel += name[index]
            tempExcel += "/"
        endExcel = tempExcel + "生成的表.xls"
        msgExcel = tempExcel + "转换信息表.xls"

        global formatErrorDir,runErrorDir
        tempStr = time.strftime('%Y-%m-%d %H-%M-%S', time.localtime(time.time()))
        formatErrorDir = tempExcel + '格式错误 ' + tempStr
        runErrorDir = tempExcel + '运行出错 ' + tempStr
        os.mkdir(formatErrorDir)
        os.mkdir(runErrorDir)
        global we, re
        re = xlrd.open_workbook(str)
        we = copy(re)

        global sheetRow, firstRow
        firstRow = 7
        sheetRow = [7, 7, 7, 7]

        global dict1,dict2,dict3,dict4,dict5
        # sheet1的字典
        dict1 = {}
        # sheet2的字典
        dict2 = {}
        # sheet3的字典
        dict3 = {}
        # sheet4的字典
        dict4 = {}
        # sheet5的字典
        dict5 = {}

        global successList,errorList,notSimpleList,runErrorList,totalList

        successList = []
        errorList = []
        notSimpleList = []
        runErrorList = []
        totalList = []

    #初始化图形界面
    def initGUI(self):
        self.progressBar.setValue(0)
        self.textEdit.setText("")

    #设置用户的操作行为
    def setOp(self, flag):
        self.startBtn.setEnabled(flag)
        self.wordDirBtn.setEnabled(flag)
        self.excelBtn.setEnabled(flag)
        self.wordTemBtn.setEnabled(flag)
        self.wordDirLineEdit.setEnabled(flag)
        self.wordTemLineEdit.setEnabled(flag)
        self.excelLineEdit.setEnabled(flag)

    #开始转换绑定的槽函数
    def startTrans(self):
        if self.wordDirLineEdit.text() == '' or self.wordTemLineEdit.text() == '' or self.excelLineEdit.text() == '':
            msgBox = QMessageBox(QMessageBox.Warning, "警告", "请选中要转换的目录")
            msgBox.exec()
            return
        self.setOp(False)
        self.statusbar.showMessage("正在进行初始化操作")
        self.initConfig()
        readTemplate(self)
        path = self.wordDirLineEdit.text().strip()
        fileTotal = 0
        total = 0
        for root, dirs, files in os.walk(path):
            fileTotal += len(files)
        print (fileTotal)
        for root, dirs, files in os.walk(path):
            self.statusbar.showMessage("正在录入(" + root + ")的文件")
            for doc in files:
                length = len(doc)
                index = doc.rfind('.')
                after = doc[index:length]
                after = after.lower()
                total += 1
                print (doc)
                print (total)
                if after != '.docx':
                    continue
                value = total * 1.0 / fileTotal * 100
                QApplication.processEvents()
                self.progressBar.setValue(int(value))
                # 是生成的临时文件，开始下一次循环
                if doc.startswith('~$'):
                    continue
                totalList.append(doc)
                try:
                    writeExcel(doc, root + '/' + doc, self)
                except Exception as e:
                    runErrorList.append(doc)
                    shutil.copyfile(root + '/' + doc, runErrorDir + '/' + doc)
                    print(e)

        writeMsg(self)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    main = SimpleDialogForm()
    main.show()
    sys.exit(app.exec_())